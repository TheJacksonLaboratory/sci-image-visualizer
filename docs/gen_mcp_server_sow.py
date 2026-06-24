"""Generate JIT_image_visualization_MCP_server_SOW.docx — SOW for adding a Model
Context Protocol (MCP) server that lets a Claude (or any MCP client) session
control the jax-image-visualization library's plotting/viewer.

Investigates jit-ui#97 ("Investigate JIT-mcp server to control the
image-visualization"). Models its structure and helpers on
gen_jit_ui_visualization_sow.py so the two SOWs read consistently."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

# Resolve outputs relative to this script's directory (libs/jax-image-visualization/docs).
_HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(_HERE, "JIT_image_visualization_MCP_server_SOW.docx")
ARCH_IMG = os.path.join(_HERE, "img", "jit-ui-visualization-architecture.png")
FIG_HEIGHT = Inches(8.6)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def num(text):
    return doc.add_paragraph(text, style="List Number")


def labeled(label, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(value)
    return p


def table(rows, header):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    return t


# ---------------- TITLE ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Image-Visualization MCP Server")
r.bold = True
r.font.size = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Statement of Work")
r.italic = True
r.font.size = Pt(14)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle2.add_run(
    "A Model Context Protocol server that lets a Claude session drive the "
    "jax-image-visualization viewer"
)
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()

labeled("Document", "JIT_image_visualization_MCP_server_SOW")
labeled("Status", "Investigation / proposal — not yet committed work. Spawned from jit-ui#97.")
labeled("Owner", "Baha Elkassaby")
labeled("Date", date.today().isoformat())
labeled("Tracking issue", "jit-ui#97 — \"Investigate JIT-mcp server to control the image-visualization\".")
labeled("Scope", "jax-image-visualization (Angular 17 / Nx buildable library) — a new opt-in control adapter and a separate MCP server package; the host app (jit-ui) wires the transport. Builds on the stable IVisualizer contract (visualization SOW D7).")
labeled("References", "mcpmarket.com/server/visualization and github.com/xlisp/visualization-mcp-server (both cited in jit-ui#97).")

doc.add_paragraph()

# ---------------- 1. EXECUTIVE SUMMARY ----------------
h1("1. Executive summary")
para(
    "This SOW investigates adding a Model Context Protocol (MCP) server so a Claude session "
    "(or any MCP client) can control the jax-image-visualization viewer programmatically — open "
    "an image, switch plot type, pan/zoom, toggle channels/colormap, create and edit regions, run "
    "browser-side SAM/Cellpose segmentation, and read back the current view or region geometry."
)
para(
    "The single most important finding, because it determines the whole design: the two reference "
    "servers cited in the issue and this library solve opposite problems. The reference servers "
    "(mcpmarket visualization and xlisp/visualization-mcp-server) are stateless plot generators — the "
    "client sends data, they return a PNG/HTML file. jax-image-visualization is a stateful, "
    "browser-resident, interactive viewer (OpenSeadragon + Plotly behind one IVisualizer contract). "
    "\"Controlling the plotting from a Claude session\" therefore means driving a live UI inside a "
    "browser, not rendering a file. That is a remote-control bridge problem, not a plot-generation "
    "problem."
)
para(
    "The library is unusually well-positioned for this because the entire UI is already driven "
    "through one backend-neutral command interface — the IVisualizer contract and its concrete "
    "RoutingVisualizerService. Every action a Claude session would want is already a typed method "
    "there. The work is therefore not to design a new control API but to expose the existing one "
    "over a transport an MCP server can reach."
)
para(
    "Recommendation: build a live-control bridge (Approach A in §3) — a thin, serializable control "
    "adapter over IVisualizer, exposed over a WebSocket the host app opens, with a separate Node MCP "
    "server package mapping one MCP tool per command. A headless variant (Approach B) reuses ~90% of "
    "the same code for automation/CI. The stateless-generator approach (Approach C) should be "
    "explicitly declined: it duplicates the rendering engine and discards everything the library is "
    "good at (tiling, regions, browser-side segmentation)."
)

# ---------------- 2. BACKGROUND: THE CONTROL SURFACE ----------------
h1("2. Background: the existing control surface")
para(
    "The library already exposes one backend-neutral command interface that the whole UI is driven "
    "through, so no new control API is needed — only a transport in front of it."
)
bullet("IVisualizer contract (contracts/visualizer.contract.ts) — ~80 methods composed from role interfaces: IDataRenderer (load/plot/zoom/stack navigation/pixel readback), IRegionStore (region CRUD, selection, undo/redo, GeoJSON import/export), IToolController (wand, brush, vertex eraser, box/point SAM, Cellpose), IDisplayOptions (colormap/LUT, reverse scale), plus capability-gated 3D / isosurface / intensity controls.")
bullet("RoutingVisualizerService — the concrete implementation bound to the VISUALIZER DI token; the composition-root facade the host already calls. Every action the MCP server would expose is already a typed method here.")
bullet("PlotType enum — IMAGE (OpenSeadragon), HEATMAP, CONTOUR, SCATTER, SURFACE, SCATTER3D, ISOSURFACE (Plotly). Plot switching is a single setPlotType() call.")
bullet("Read-back paths already exist — getRegions()/getGeoJsonString() for geometry, getDisplayedPixelData()/getCurrentImage() for pixels, and downloadImage()/exportComposite() for a rendered PNG the client can be handed as MCP image content.")

para(
    "The architectural gap is purely one of transport and process boundary: an MCP server is a "
    "Node/Python process, while the visualizer lives in a browser tab. The whole design question is "
    "what bridges the two."
)
para(
    "Claude session  --(MCP stdio/SSE)-->  MCP server (Node process)  --( ??? )-->  "
    "RoutingVisualizerService (Angular, in a browser tab)",
    italic=True,
)

# Context: reuse the existing architecture diagram to show where the bridge attaches.
doc.add_page_break()
doc.add_picture(ARCH_IMG, height=FIG_HEIGHT)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
_cap = doc.add_paragraph()
_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
_cap.add_run(
    "Figure 1 — jax-image-visualization architecture. The MCP control adapter attaches at the "
    "RoutingVisualizerService seam (top of the library), reusing the same IVisualizer commands the "
    "toolbar and host already drive."
).italic = True
doc.add_page_break()

# ---------------- 3. APPROACHES CONSIDERED ----------------
h1("3. Approaches considered")
para(
    "Three approaches are genuinely different products. They differ chiefly in who owns the browser "
    "the viewer runs in."
)

h2("3.1 Approach A — Live-control bridge (recommended)")
para(
    "The MCP server talks over a WebSocket to a thin control adapter wired to the "
    "RoutingVisualizerService in the running app — the same tab the scientist is looking at. Claude "
    "drives the real UI; results are read back as region GeoJSON or as a screenshot via "
    "downloadImage()/exportComposite() returned over the channel."
)
labeled("Controls the real UI the analyst sees", "Yes — same tab.")
labeled("Reuses IVisualizer as-is", "Yes — fully.")
labeled("OSD tiles / SAM / regions", "Yes — the host supplies its existing DI ports.")
labeled("Effort", "Medium.")
labeled("Best for", "\"Segment the cells in the rectangle and switch to contour view\" on the analyst's actual screen.")

h2("3.2 Approach B — Headless / embedded renderer")
para(
    "The MCP server launches the library in a headless browser (e.g. Playwright) it owns, drives it, "
    "and screenshots back natively. Same control adapter as A; the difference is the server owns the "
    "browser instead of attaching to a human's tab. Best for reproducible/automated/CI visualization "
    "with no human tab open. Requires the planned example server (visualization SOW D7) to stand in "
    "for jit-service so the ports resolve standalone."
)
labeled("Effort", "Medium-High (adds headless-browser orchestration on top of A's adapter).")

h2("3.3 Approach C — Stateless generator (declined)")
para(
    "Port the Plotly trace-building to Node and render data to a PNG, mirroring the reference "
    "servers. Declined: it reimplements the rendering engine and throws away tiling, regions, and "
    "browser-side SAM/Cellpose — i.e. everything the library is good at. Recorded here so the issue "
    "can close the option explicitly rather than leave it open."
)

h2("3.4 Comparison")
table([
    ["Controls the real UI the analyst sees", "Yes (same tab)", "A separate headless instance", "No (throwaway plots)"],
    ["Claude can \"see\" the result", "Screenshot over the channel", "Native (Playwright)", "Returns the image"],
    ["Reuses IVisualizer as-is", "Yes", "Yes", "No (reimplements rendering)"],
    ["OSD tiles / SAM / regions work", "Yes", "Yes (ports stubbed)", "No"],
    ["Effort", "Medium", "Medium-High", "High + duplicates logic"],
    ["Recommendation", "Recommended", "Recommended (headless variant)", "Declined"],
], header=["Dimension", "A. Live-control bridge", "B. Headless renderer", "C. Stateless generator"])
para(
    "A and B share roughly 90% of the code — the MCP tool layer and the control adapter — and differ "
    "only in who owns the browser. Building A first yields B for little extra.",
    italic=True,
)

# ---------------- 4. PROPOSED ARCHITECTURE ----------------
h1("4. Proposed architecture (Approach A)")
num("Control adapter in the library (src/lib/mcp/visualizer-control.adapter.ts): an injectable that takes VISUALIZER and exposes a flat, serializable command set (loadImage, setPlotType, setColormap, setRegionsGeoJson, segmentRectangles, zoomIn/zoomOut, getScreenshotPng, getRegionsGeoJson, ...). A thin JSON-in/JSON-out translation of IVisualizer. This is the only new code inside the library, and it lives behind a new opt-in entry point so the core library stays transport-free.")
num("Transport: a WebSocket client the host app opens (dev / feature-flagged only) that registers the adapter's commands and relays calls/results.")
num("MCP server (new package, e.g. tools/jit-viz-mcp/, Node + @modelcontextprotocol/sdk): one MCP tool per adapter command; the server relays over the WebSocket and returns results and screenshots as MCP content. This is the package added to .mcp.json (which already exists in the repo as a scaffold, currently only generic github/filesystem/git servers, all commented out).")
para(
    "This lines up with the \"example/test server + demos\" task already on the library roadmap "
    "(README \"In progress / roadmap\" and the extraction SOW). That standalone example server is the "
    "natural host for the bridge in headless mode (Approach B), so the MCP work can ride on "
    "already-planned infrastructure rather than depending on jit-ui being open."
)

# ---------------- 5. DELIVERABLES ----------------
h1("5. Deliverables")
para("All deliverables are Proposed (none started). This SOW is an investigation outcome, not committed work.")

h2("M1. Control adapter over IVisualizer")
labeled("What", "A serializable, transport-neutral command adapter that wraps RoutingVisualizerService / IVisualizer: load image, set plot type, set colormap/reverse scale, zoom/pan, stack navigation, region get/set as GeoJSON, run box/point SAM and Cellpose, and read back regions and a screenshot. Behind a new opt-in library entry point so the core library carries no transport dependency.")
labeled("Where", "libs/jax-image-visualization/src/lib/mcp/* + a secondary entry point in the barrel.")
labeled("Status", "Proposed.")

h2("M2. WebSocket control transport + host wiring")
labeled("What", "A WebSocket channel the host app opens (dev / feature-flagged) that registers the M1 commands, plus a session/registration handshake identifying which viewer tab is being controlled.")
labeled("Where", "library transport module + apps/jit-ui composition root (flagged).")
labeled("Status", "Proposed.")

h2("M3. MCP server package")
labeled("What", "A Node package built on @modelcontextprotocol/sdk exposing one MCP tool per M1 command, relaying over the M2 transport, and returning screenshots as MCP image content. Registered in .mcp.json.")
labeled("Where", "tools/jit-viz-mcp/ (new package) + .mcp.json.")
labeled("Status", "Proposed.")

h2("M4. Headless mode (optional, Approach B)")
labeled("What", "Run the library in a Playwright-driven headless browser the MCP server owns, against the planned example server, for automation/CI with no human tab. Reuses M1/M3; adds browser orchestration and example-server data stubs.")
labeled("Where", "tools/jit-viz-mcp/ headless driver + the library example server (visualization SOW D7).")
labeled("Status", "Proposed (depends on the example server).")

h2("M5. Design doc")
labeled("What", "A docs/mcp-control-design.md modeled on docs/sam-segmentation-design.md: the command catalogue, transport/handshake, security model, and the serializable subset of the IVisualizer surface.")
labeled("Where", "libs/jax-image-visualization/docs/mcp-control-design.md.")
labeled("Status", "Proposed.")

# ---------------- 6. MVP ----------------
h1("6. Suggested MVP")
para(
    "A compelling first cut is read state plus five commands: load, setPlotType, setColormap, "
    "setRegionsGeoJson, and getScreenshot. That alone demonstrates \"Claude drives the viewer\" and "
    "validates the transport and handshake before building out all ~80 IVisualizer methods. Expand "
    "the command catalogue (segmentation, tools, 3D controls) only once the transport is proven."
)

# ---------------- 7. SECURITY & OPEN QUESTIONS ----------------
h1("7. Security and open questions")
num("Auth — both deployments sit behind oauth2-proxy + Auth0. A control channel that can run segmentation and export data needs the same auth story. Simplest first step: gate the channel to dev/local behind a feature flag.")
num("Single-tab vs multi-session — the bridge controls whichever tab registered; a session/registration handshake is required once more than one viewer can be open.")
num("any-typed methods — the IVisualizer contract is deliberately permissive (any on several Plotly methods). The adapter is a good forcing function to tighten the serializable subset; decide how strict M1's schema must be.")
num("Transport choice — WebSocket (live tab) vs an SSE/HTTP control endpoint vs headless-only. Affects how the host registers and how auth is enforced.")
num("Where the MCP package lives — in this repo (tools/) vs a sibling repo, and how it is versioned against the library's public contract.")

# ---------------- 8. DEPENDENCIES ----------------
h1("8. Dependencies")
bullet("A stable IVisualizer / RoutingVisualizerService public contract (visualization SOW D7 — library extraction hardening). The adapter binds to this surface.")
bullet("@modelcontextprotocol/sdk (Node) for the MCP server; a WebSocket library for the transport.")
bullet("For headless mode (M4): the planned library example server (visualization SOW D7) and Playwright (or equivalent) for the headless browser.")
bullet("The host app's oauth2-proxy + Auth0 setup for any non-dev deployment of the control channel.")

# ---------------- 9. RISKS ----------------
h1("9. Risks")
table([
    ["Control channel becomes an unauthenticated remote-control surface",
     "Medium", "High",
     "Gate to dev/local behind a feature flag initially; align the channel with oauth2-proxy + Auth0 before any shared deployment; require a session handshake."],
    ["IVisualizer's any-typed methods leak weak schemas into MCP tools",
     "Medium", "Low",
     "Define a tightened, serializable command subset in M1; validate at the adapter boundary rather than exposing raw methods."],
    ["Drift between the MCP command catalogue and the evolving library contract",
     "Medium", "Medium",
     "Generate the command list from the contract where possible; version the MCP package against the library's public API; keep the MVP surface small."],
    ["Headless mode blocked by the example server not yet existing",
     "Medium", "Low",
     "Ship Approach A (live tab) first; treat M4 as dependent on visualization SOW D7."],
    ["Single-tab ambiguity when multiple viewers are open",
     "Low", "Medium",
     "Registration/handshake that names the controlled tab; reject commands when the target is ambiguous."],
], header=["Risk", "Likelihood", "Impact", "Mitigation"])

# ---------------- 10. PHASING ----------------
h1("10. Phasing and acceptance criteria")
h2("10.1 Phase 1 — MVP live-control bridge (M1 partial, M2, M3 partial)")
labeled("Acceptance", "From a Claude session, load an image, switch plot type, set a colormap, push regions as GeoJSON, and get a screenshot back — driving a real jit-ui tab over the WebSocket channel (dev/flagged).")
h2("10.2 Phase 2 — Full command catalogue (M1, M3)")
labeled("Acceptance", "Zoom/pan, stack navigation, region read-back, box/point SAM and Cellpose, and 3D/isosurface controls are all exposed as MCP tools with validated schemas; read-back returns region geometry and pixel/screenshot content.")
h2("10.3 Phase 3 — Security hardening")
labeled("Acceptance", "The control channel authenticates via oauth2-proxy + Auth0 (or a documented equivalent) and enforces a session handshake; the dev-only flag is no longer required for a controlled deployment.")
h2("10.4 Phase 4 — Headless mode (M4, optional)")
labeled("Acceptance", "The MCP server drives a headless browser against the example server and returns screenshots with no human tab open; runnable in CI.")

# ---------------- 11. CHANGELOG ----------------
h1("11. Changelog")
labeled(date.today().isoformat(), "Initial draft. Investigation outcome for jit-ui#97: documents the control-surface finding (IVisualizer / RoutingVisualizerService already expose every needed command), the three approaches (live-control bridge recommended, headless variant, stateless generator declined), the proposed architecture, deliverables M1-M5, a five-command MVP, security/open questions, dependencies, risks, and phasing.")

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
