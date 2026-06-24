"""Generate JIT_UI_visualization_library_SOW.docx — SOW for the jit-ui
visualization rework and its extraction into the jax-image-visualization library."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date

# Resolve outputs relative to this script's directory (libs/jax-image-visualization/docs)
# so the generator is portable — it no longer hardcodes a jit-service path.
_HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(_HERE, "JIT_UI_visualization_library_SOW.docx")
ARCH_IMG = os.path.join(_HERE, "img", "jit-ui-visualization-architecture.png")
REGION_IMG = os.path.join(_HERE, "img", "jit-ui-region-architecture.png")
# Section 9 "Current visuals" — product screenshots, sized by width to the page.
VIS_DIR = os.path.join(_HERE, "img", "sow-visuals")
VIS_WIDTH = Inches(6.5)
# Both diagrams are tall (portrait) ELK layouts. Each gets its own page (page
# breaks around it) and is sized by height to fill that page — leaving just room
# for the caption under it.
FIG_HEIGHT = Inches(8.6)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def num(text):
    return doc.add_paragraph(text, style="List Number")


def labeled(label, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(value)
    return p


def table(rows, header):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    return t


# ---------------- TITLE ----------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("jit-ui Visualization Rework and Library Extraction")
r.bold = True
r.font.size = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Statement of Work")
r.italic = True
r.font.size = Pt(14)

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle2.add_run("A backend-neutral, reusable image-visualization library extracted from jit-ui")
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()

labeled("Document", "JIT_UI_visualization_library_SOW")
labeled("Status", "Substantially implemented and merged to jit-ui master (via the 'openseadragon' branch, PR #79); library extraction in progress, distribution pending")
labeled("Owner", "Baha Elkassaby")
labeled("Date", date.today().isoformat())
labeled("Scope", "jit-ui (Angular 17 / Nx) — the visualization layer and the new buildable library 'jax-image-visualization'; backed by jit-service tile/histogram/export endpoints")
labeled("Realized library name", "jax-image-visualization (Nx buildable library)")
labeled("Distribution", "npm under the @jax-data-science organization scope (Nx publishable library); package name TBD among image-visualization, jax-image-visualization, or jit-image-visualization")

doc.add_paragraph()

# ---------------- 1. EXECUTIVE SUMMARY ----------------
h1("1. Executive summary")
para(
    "This SOW covers the rework of the jit-ui visualization layer and its extraction into a standalone, "
    "backend-neutral Angular library, jax-image-visualization."
)
para(
    "The visualization layer was redesigned around a single backend-neutral contract with two pluggable "
    "rendering backends: a natively-tiled OpenSeadragon viewer for the deeply-zoomable Image plot type "
    "(backed by the jit-service tile endpoints), and Plotly for the scientific plot types (heatmap, contour, "
    "scatter, surface, scatter-3D, isosurface). A backend-neutral region model and on-canvas tools work "
    "identically on either backend, and a Channels and Histogram tool provides per-channel "
    "brightness/contrast, pseudo-colour compositing, and true 16-bit support."
)
para(
    "Most of the engineering (the contract, both backends, regions/tools, the plot-type framework, and the "
    "Channels and Histogram tool including 16-bit) is already implemented on jit-ui master (merged from the "
    "'openseadragon' branch) and the supporting jit-service endpoints are deployed. The remaining work is finishing the "
    "library extraction (clean ports/DI boundaries so the package carries no host coupling) and deciding "
    "and executing distribution so other JAX Data Science frontends can consume it. Each deliverable in §4 "
    "carries an explicit status (Implemented / In progress / Planned)."
)

h2("1.1 Current architecture")
para(
    "The library is organized around one backend-neutral contract and a router that dispatches to a "
    "rendering backend per plot type. The host app embeds a single <visualization> component (which wraps "
    "the toolbar and the visualizer); the Channels and Histogram dialog and the region editor are sibling "
    "embeddable components. Everything app-specific is reached through dependency-injection ports the host "
    "supplies, so the library depends on no host code."
)
doc.add_page_break()
doc.add_picture(ARCH_IMG, height=FIG_HEIGHT)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
_cap = doc.add_paragraph()
_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
_cr = _cap.add_run("Figure 1 — jit-ui visualization architecture (jax-image-visualization).")
_cr.italic = True
bullet("Host app: the diagram component embeds the <visualization> component, which wraps the plotting-toolbar (Toolbar) and the visualizer (#plot, where the active backend renders). The Channels and Histogram dialog and the region editor are sibling embeddable components.")
bullet("Public API: IVisualizer (composed of IDataRenderer, IRegionStore, IToolController and IDisplayOptions, plus a ViewerCapabilities descriptor), IChannelHistogramApi (CHANNEL_HISTOGRAM_API) and IRegionEditorApi — plus the DI ports TILE_ACCESS_PORT, IMAGE_STATE_PORT, REGION_IO_PORT and VIZ_CONFIG. Components consume the tokens; the host provides the ports (see 1.2).")
bullet("Router: RoutingVisualizerService implements the three API tokens (the host binds them useExisting), injects the ports, and routes the Image plot type to OpenSeadragon and the scientific plot types to Plotly.")
bullet("Backends and shared state: the OpenSeadragon and Plotly backends both implement IVisualizer and both read the shared VisualizerStore (display/channel/colormap state) and RegionStore (regions, selection, class colours).")
bullet("Backend internals: the backend-neutral region overlay and on-canvas tools (wand, vertex-eraser, zoom-to-box, rectangle/freehand/polyline, bezier) and the OpenSeadragon tile pipeline (recolor LUT, scale bar, slice cache with background loading).")
bullet("Server: both backends fetch image data from the jit-service endpoints — /preview (the downscaled overview / untiled image, URL supplied via IMAGE_STATE_PORT), /tiles/info and /tile (per-channel tiles, OpenSeadragon), /zoom/region (Plotly's high-resolution re-fetch via TILE_ACCESS_PORT.zoomOnRegion), /histogram and /export/tiff.")

para(
    "Regions are created and edited through their own small set of contracts, so the tools and overlays stay "
    "decoupled from the backends and from the region state. The toolbar routes each tool mode either to the "
    "per-backend region overlay (IRegionOverlay — an OpenSeadragon SVG overlay and a thin Plotly native-shapes "
    "overlay) or to an on-canvas tool through IToolController. The wand, brush and vertex-eraser tools bind to a "
    "small host interface (IViewportHost + IRegionDataHost) and share the WandService geometry helpers; every "
    "overlay and tool reads and writes the single shared RegionStore (IRegionStore + IRegionEditApi)."
)
doc.add_page_break()
doc.add_picture(REGION_IMG, height=FIG_HEIGHT)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
_cap2 = doc.add_paragraph()
_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
_cr2 = _cap2.add_run("Figure 2 — region interfaces and region tools (jax-image-visualization).")
_cr2.italic = True

h2("1.2 Host integration: DI ports and adapters")
para(
    "The library is host-agnostic: everything app-specific is expressed as a port (a TypeScript interface "
    "plus an Angular InjectionToken). The host application's composition root provides one adapter per port "
    "and binds the API tokens to the library's router. This is the seam that lets the package carry no host "
    "coupling and be published and reused (deliverables D7 and D8)."
)
table([
    ["VISUALIZER · CHANNEL_HISTOGRAM_API · REGION_EDITOR_API", "RoutingVisualizerService (useExisting)", "The render/viewport, channel-histogram and region-editor APIs the embeddable components call. One library service implements all three."],
    ["TILE_ACCESS_PORT", "TileAccessAdapter (host)", "Selected image's base64 FileInfo, bearer auth headers, ROI-zoom and diagram-display selection. The OpenSeadragon backend uses it (with the VIZ_CONFIG base URL) to fetch tiles, histograms and exports from jit-service."],
    ["IMAGE_STATE_PORT", "ImageStateAdapter (host)", "Current image info, filename, loading and cache-progress, zoom flag and panel width (read + write). Bridges the viewer to the app's image/loading state so the right overlays and progress show."],
    ["REGION_IO_PORT", "RegionIoAdapter (host)", "Selected filename, ROI-file existence check and GeoJSON save. Bridges region import/export persistence to the app's storage."],
    ["VIZ_CONFIG", "useValue (environment)", "Backend base URL (slideCropServer) and optional UI hints (e.g. regionEditorWidthSelector)."],
], header=["Token / port", "Provided by", "Responsibility"])
para(
    "Because every host touch-point is one of these adapters, the library imports nothing from jit-ui. A "
    "downstream JAX Data Science frontend reuses it by providing its own four adapters plus a VIZ_CONFIG "
    "value (and the useExisting bindings to RoutingVisualizerService, or its own IVisualizer "
    "implementation) — with no library source changes."
)

h2("1.3 Consuming the library: server-backed and serverless")
para(
    "An external app embeds <visualization> (and optionally <region-editor> / <channel-histogram>) and "
    "provides the four DI ports plus VIZ_CONFIG; the library imports nothing from the host. One field on the "
    "image descriptor the host streams from IMAGE_STATE_PORT — IImageInfo.tiled — selects how pixels are "
    "sourced, so the same component supports two consumption modes with no library changes."
)
para("Mode A — server-backed (a jit-service-like backend that returns URLs).", bold=True)
bullet("The host's IMAGE_STATE_PORT emits IImageInfo with tiled: true (the default) and urls[zIndex] pointing at the backend's downscaled overview (/preview); optionally smallUrls[] (fast placeholder) and mppX/Y (scale bar).")
bullet("OpenSeadragon drives its tile pipeline: it polls /tiles/info and fetches /tile off VIZ_CONFIG.slideCropServer, using the base64 info from TILE_ACCESS_PORT.getSelectedInfoB64() and TILE_ACCESS_PORT.getAuthHeaders(); Plotly's high-resolution re-fetch on zoom goes through TILE_ACCESS_PORT.zoomOnRegion() (/zoom/region). Region persistence goes through REGION_IO_PORT.saveGeoJson().")
bullet("The host builds a backend exposing (at minimum) /preview, /tiles/info, /tile — plus /zoom/region, /histogram and /export/tiff for Plotly hi-def zoom and true 16-bit. Endpoint shapes are the host's choice; the library only consumes the URL string and the port methods. This is the path for whole-slide / deeply-zoomable / multi-channel 16-bit imagery. (jit-ui uses this mode.)")
para("Mode B — serverless (in-memory pixels, no backend).", bold=True)
bullet("The host turns already-decoded pixels (canvas / ImageData) into a blob: or data: URL and emits IImageInfo with tiled: false and urls: [thatUrl], sized to the image's own pixels.")
bullet("Because tiled is false, OpenSeadragon opens that single self-contained image directly — no /tiles/info, no /tile, no slideCropServer. The remaining ports are no-ops (TILE_ACCESS_PORT returns null/EMPTY/{}, REGION_IO_PORT returns false/void), and VIZ_CONFIG.slideCropServer is empty.")
bullet("The host builds essentially no server code — just an IMAGE_STATE_PORT adapter that emits the blob/data URL with tiled: false, plus stub adapters. This is the path for viewport-sized, already-decoded images (e.g. the in-app image-processing-pipeline preview). Trade-off: no deep tiled zoom, no server-side hi-def re-fetch, and 8-bit display fidelity.")
para(
    "The decision is a single switch on the streamed descriptor — IImageInfo.tiled === false routes to the "
    "direct single-image load (Mode B); otherwise OpenSeadragon tiles via the server (Mode A). Both modes use "
    "the same component, regions/tools, and Channels and Histogram machinery. D7's example server + examples "
    "ship a runnable reference for each path."
)

# ---------------- 2. ORIGINALITY AND PRIOR ART ----------------
h1("2. Originality and prior art")
para(
    "This section situates the library against the existing landscape of scientific-image tools — desktop "
    "applications, server-backed web platforms, and embeddable JavaScript libraries — and states precisely what "
    "is original here and what is not. The short version: every individual capability below exists somewhere in "
    "a mature incumbent, but no embeddable, backend-neutral web library unifies all of them behind one contract, "
    "and no production bioimage tool runs domain-finetuned ViT segmentation fully client-side. The originality is "
    "in the integration and the delivery model, not in claiming any single capability for the first time."
)

h2("2.1 Landscape at a glance")
para(
    "The table compares representative tools across the four axes this library spans — deep-zoom whole-slide "
    "(WSI) tiling, scientific plotting (heatmap / contour / surface / isosurface), annotation/region tooling, "
    "per-channel and true 16-bit display — plus where AI segmentation runs and whether the tool is embeddable "
    "as a library (versus a standalone application or a server-coupled platform). The final row is this library."
)
table([
    ["Fiji / ImageJ (desktop)", "Via Bio-Formats", "Limited", "Rich (ROI, wand)", "True 8/16/32-bit", "Local plugin (CPU)", "No"],
    ["napari (desktop)", "Multiscale (dask/zarr)", "Limited", "Shapes / labels", "Yes (n-D, 16-bit)", "Desktop GPU plugin", "No"],
    ["QuPath (desktop)", "Yes (native)", "No", "Rich, GeoJSON", "Yes (multiplex)", "Desktop / SAM-API server", "No"],
    ["Imaris (desktop, commercial)", "No (volume .ims)", "3D volume render", "Surfaces / spots", "Yes", "Local GPU", "No"],
    ["3D Slicer / ITK-SNAP (desktop)", "No (medical volumes)", "3D volume render", "Segment editor", "Yes (float)", "Local GPU (MedSAM/MONAI)", "No"],
    ["OMERO.web + iviewer (web)", "Yes (server)", "No", "Yes", "Server-rendered 16-bit", "Server-side / none", "No (platform)"],
    ["Cytomine / CaMicroscope (web)", "Yes", "No", "Rich", "Partial", "Server-side", "No (platform)"],
    ["Viv / Avivator (web library)", "Yes (OME-TIFF/Zarr)", "No", "Basic", "Yes (GPU, native 16-bit)", "No", "Library"],
    ["Vitessce (web framework)", "Yes (via Viv)", "Scatter / heatmap", "Cell / region", "Yes (via Viv)", "No", "React component"],
    ["Cornerstone3D / OHIF (web, DICOM)", "Yes (DICOM)", "No", "Measurement tools", "Yes", "No (server)", "Library / app"],
    ["OpenSeadragon (web library)", "Yes", "No", "Plugins only", "No", "No", "Library"],
    ["Plotly.js (web library)", "No", "Yes", "No", "No", "No", "Library"],
    ["Annotorious (web library)", "Via OSD", "No", "Yes (W3C)", "No", "No", "Library"],
    ["jax-image-visualization (this SOW)", "Yes (OpenSeadragon)", "Yes (Plotly)", "Rich, GeoJSON", "Yes (true 16-bit)", "In-browser ViT (ONNX)", "Library"],
], header=["Tool / family", "WSI tiling", "Sci. plots", "Annotation", "Channels / 16-bit", "AI segmentation", "Embeddable?"])
para(
    "Read across the rows: the established tools each leave at least one axis empty, and they cluster into three "
    "groups — single-purpose web libraries (OpenSeadragon, Plotly, Annotorious, Viv), full desktop applications "
    "(Fiji, napari, QuPath, Imaris, Slicer), and server-coupled web platforms (OMERO, Cytomine, CaMicroscope). "
    "Only the final row fills every column from a single embeddable component.",
    italic=True,
)

h2("2.2 Desktop scientific-imaging tools")
para(
    "Fiji/ImageJ, napari, QuPath, Imaris, Aivia, Icy, 3D Slicer and ITK-SNAP are the entrenched, well-cited "
    "tools of scientific image analysis, and on raw analytical depth several of them exceed what a browser "
    "library can do: Imaris and 3D Slicer set the bar for GPU volume ray-casting and 4D tracking; napari offers "
    "native n-dimensional rendering; QuPath provides whole-slide cell detection, classification and "
    "quantification pipelines this library does not attempt to match. The differentiator is not capability on "
    "those axes — it is the distribution model. Every one of these tools requires a local install (a Java "
    "runtime, a Python/conda/CUDA environment, or a commercial Windows/macOS binary), and none ships as a web "
    "component or npm package that drops into a web application. Their AI segmentation (SAMJ, micro-sam, "
    "QuPath's Cellpose/StarDist/SAM-API, Slicer's MedSAM/MONAI) runs on a local desktop GPU/CPU or a separately "
    "hosted Python server — never quantized client-side in the browser."
)

h2("2.3 Web viewers and WSI platforms")
para(
    "On the web, OpenSeadragon is the de-facto deep-zoom tiled viewer, but by itself it is only that: no "
    "scientific plots, no built-in region model, no per-channel or 16-bit handling, no segmentation (annotation "
    "comes from bolt-on plugins such as Annotorious). The full WSI platforms — OMERO.web + omero-iviewer, "
    "Cytomine, CaMicroscope, Orbit, Pathomation — bundle viewing, annotation and (server-side) analysis, but "
    "they are standalone, backend-heavy applications, not embeddable libraries, and their AI segmentation runs "
    "on the server. The closest web peers on the imaging axis are HMS-DBMI's Viv (GPU-accelerated multichannel "
    "and native 16-bit rendering of OME-TIFF/Zarr via deck.gl) and Vitessce (which builds linked single-cell "
    "dashboards on Viv). Viv is excellent at multichannel display — and exceeds this library there — but it does "
    "not pair a tiled viewer with a scientific-plotting engine behind one API, has only minimal annotation, and "
    "does no in-browser segmentation; Vitessce is React-coupled, locked to the single-cell/spatial domain, and "
    "likewise has no client-side ViT segmentation."
)

h2("2.4 Embeddable JavaScript libraries and the architecture")
para(
    "Among embeddable libraries the pattern is consistently single-purpose: Plotly.js renders plots, "
    "OpenSeadragon renders tiles, Annotorious renders annotations, Viv renders multichannel pixels, "
    "Cornerstone3D/OHIF render DICOM. Cornerstone3D is the strongest precedent for a layered, embeddable "
    "imaging toolkit with host-supplied loaders and metadata providers — conceptually similar to this "
    "library's data ports — but it is DICOM/radiology-centric, with no WSI deep-zoom focus, no Plotly-class "
    "scientific plotting and no in-browser ViT segmentation. What is architecturally uncommon here is the "
    "backend-neutral contract itself: a router (RoutingVisualizerService) that dispatches by plot type to "
    "pluggable rendering backends (OpenSeadragon for the Image type, Plotly for the scientific types) while the "
    "host injects pluggable data backends through DI ports (server-tiled, §1.3 Mode A, or serverless in-memory, "
    "Mode B). This hexagonal port/adapter design is routine in application architecture but rare as the public "
    "API of a visualization library — and it is what lets the same component serve a whole-slide server-backed "
    "viewer and an in-app, in-memory pipeline preview with no source changes."
)

h2("2.5 In-browser ViT segmentation")
para(
    "In-browser inference of the original Segment Anything Model is, by itself, no longer novel: generic-model "
    "demos run fully client-side today (Transformers.js / SlimSAM, lucasgelfond's WebGPU SAM2, Microsoft's ONNX "
    "Runtime Web sample), and WebGPU is now production-viable across Chrome, Edge, Firefox and Safari with a "
    "mandatory WASM fallback. The honest claim is narrower and stronger: no production bioimage tool runs "
    "domain-finetuned biomedical ViT segmentation fully in the browser. micro-sam (microscopy), patho-sam "
    "(histopathology) and Cellpose-SAM are Python/desktop/napari or server-GPU tools with no official browser "
    "ONNX export — community attempts to convert micro-sam to ONNX have hit errors — and even the Cellpose-SAM "
    "“no-install” cloud option is a server-GPU Hugging Face Space, not client-side. Comparable bioimage "
    "tools either run the model server-side (Cytomine, DeepCell/Kiosk, QuPath's SAM-API, the Cellpose-SAM HF "
    "Space) or on a desktop GPU with Python (QuPath+Cellpose, napari+micro-sam); even a leading commercial "
    "annotator (Labelbox) keeps the heavy ViT encoder server-side and runs only the light decoder in-browser. "
    "Exporting microscopy- and pathology-finetuned ViT models (micro-sam ViT-T/ViT-B, patho-sam ViT-B + int8) "
    "and Cellpose-SAM to quantized ONNX and running them client-side via onnxruntime-web (WebGPU + WASM "
    "fallback), in both promptable and automatic modes, embedded directly in the annotation viewer (D6), removes "
    "the install/GPU/server/data-upload barrier and has no direct production precedent."
)

h2("2.6 What is original — and what is not")
para(
    "Stated plainly so the claim survives scrutiny:"
)
bullet("Original: unifying deep-zoom WSI tiling, scientific plotting, rich GeoJSON-interoperable annotation, per-channel + true 16-bit display, and in-browser ViT segmentation in one embeddable, backend-neutral web library — a combination no existing desktop tool, web platform, or JS library provides from a single component.")
bullet("Original: a backend-neutral contract with pluggable rendering backends (OpenSeadragon / Plotly) and host-injected pluggable data backends (server-tiled / serverless) — a port/adapter architecture rare as a visualization library's public API.")
bullet("Original: zero-install, zero-server, client-side ViT segmentation using domain-finetuned biomedical models (micro-sam, patho-sam, Cellpose-SAM) exported to quantized ONNX — no production bioimage tool does this in the browser.")
bullet("Not claimed as original: in-browser SAM as a technique (generic-model demos predate this), GPU multichannel/16-bit rendering (Viv leads here), W3C/GeoJSON region I/O (Annotorious), deep 3D volume rendering and whole-slide analysis pipelines (Imaris, 3D Slicer, QuPath remain more capable on those axes). Client-side inference also trades raw throughput and precision (no server GPU; SAM's fixed 1024×1024 input) for accessibility.")
para(
    "The value proposition is therefore accessibility, embeddability and web-native integration with client-side "
    "AI — not displacing the deep analytical or volumetric capabilities of specialized desktop software.",
    italic=True,
)

# ---------------- 3. SCOPE ----------------
h1("3. Scope")
h2("3.1 In scope")
bullet("A backend-neutral visualization contract (IVisualizer) composed of role interfaces (render/viewport, region store, on-canvas tools, display options), with capability gating so consumers only call what the active backend supports.")
bullet("OpenSeadragon backend for the Image plot type: tiled zoomable raster off the jit-service tile endpoints, navigator minimap, physical-units scale bar, stack-slice cache with background preloading, and a client-side recolor pipeline.")
bullet("Plotly backend for the scientific plot types and the plot-type selector that routes between backends per plot type.")
bullet("Backend-neutral region model and shared region store; region create/edit/select/delete; rectangle/freehand/polyline drawing; magic-wand, vertex-eraser and zoom-to-box tools; bezier regions; QuPath-compatible GeoJSON import/export; the region editor with editable classification colours.")
bullet("Channels and Histogram tool: per-channel display window, gamma, visibility, Fiji-style per-channel pseudo-colour compositing, live histogram, publication composite PNG export, and true 16-bit support (native-bit-depth histogram and windowing, plus a data-preserving multi-band 16-bit TIFF export).")
bullet("In-browser SAM segmentation toolset: promptable box and point Segment-Anything tools and an automatic Cellpose-SAM tool, running quantized ONNX models client-side (onnxruntime-web, WebGPU with a WASM fallback) through a pluggable model registry (micro-sam ViT-T / ViT-B, patho-sam ViT-B and int8) that ships hosted models as defaults.")
bullet("Extraction of all of the above into the Nx buildable library jax-image-visualization: library-owned data interfaces, contracts and DI ports/tokens, ng-packagr build, Jest unit tests, and host (jit-ui) wiring that provides the ports.")

h2("3.2 Out of scope")
bullet("The jit-service tile/histogram/export endpoints themselves (/tiles/info, /tile, /histogram, /export/tiff and the per-channel + 16-bit backend). They back this library and are tracked as their own jit-service work (issue #76 and the 16-bit support docs); this SOW depends on them but does not own them.")
bullet("Migrating other JAX Data Science frontends onto the library (downstream adoption is named as a Planned deliverable, D8, but the actual per-consumer integration is owned by each consumer).")
bullet("Server-side rendering or a non-Angular distribution. The library targets Angular consumers.")
bullet("Replacing Plotly or OpenSeadragon with a different rendering engine.")

# ---------------- 4. DELIVERABLES ----------------
h1("4. Deliverables")
para(
    "Each deliverable identifies what is built, where it lives, and its current status. 'Implemented' means it "
    "is built, verified, and merged to jit-ui master."
)

h2("D1. Backend-neutral visualization contract and router")
labeled("What", "A single IVisualizer contract composed of role interfaces (data renderer/viewport, region store, on-canvas tools, display options), with a capability descriptor so consumers gate on advertised features rather than backend identity. A RoutingVisualizerService routes the Image plot type to OpenSeadragon and every other plot type to Plotly, keeping shared state (regions, display options) in one place so both backends stay in sync.")
labeled("Where", "libs/jax-image-visualization/src/lib/contracts/visualizer.contract.ts, capabilities.contract.ts, routing-visualizer.service.ts, visualizer-store.service.ts.")
labeled("Status", "Implemented.")

h2("D2. OpenSeadragon image backend + server tile integration")
labeled("What", "A natively-tiled, deeply zoomable raster backend for the Image plot type, backed by the jit-service tile endpoints, with a navigator minimap, a physical-units scale bar, click-to-zoom, a stack-slice cache with LRU eviction and background preloading for flicker-free z-scrubbing, and a client-side recolor pipeline (tile-invalidated) that applies the grayscale colormap LUT and per-channel display settings.")
labeled("Where", "libs/jax-image-visualization/src/lib/implementations/osd/* (visualizer service, region overlay, scale bar, coordinate transform).")
labeled("Status", "Implemented.")

h2("D3. Backend-neutral regions, on-canvas tools, and GeoJSON I/O")
labeled("What", "A backend-neutral Region model and a shared region store both backends render from; region create/edit/select/delete; rectangle, freehand and polyline drawing; magic-wand, vertex-eraser and zoom-to-box tools implemented against a shared coordinate-transform abstraction; bezier regions; QuPath-compatible GeoJSON import/export; and a region editor that edits classification colours for every class from one dialog.")
labeled("Where", "libs/jax-image-visualization/src/lib/models/region.ts, region-store.service.ts, toolbar/* (wand, vertex-eraser, zoom-to-box), region-editor/*, and the per-backend overlays.")
labeled("Status", "Implemented.")

h2("D4. Plot-type framework (Plotly backends, capability gating, intensity profiles)")
labeled("What", "A plot-type selector exposing Image (OpenSeadragon), Heatmap, Contour, Scatter, Surface (3D), Scatter 3D and Isosurface (3D), with scalar/3D types gated to the images that support them; isosurface rendering of grayscale z-stacks with an intensity-band slider; and intensity-profile line ROIs with a live, zoom-aware inset chart.")
labeled("Where", "libs/jax-image-visualization/src/lib/contracts/plot-type.ts, implementations/plotly/* (service + intensity profile), capability gating in capabilities.contract.ts.")
labeled("Status", "Implemented.")

h2("D5. Channels and Histogram tool, incl. true 16-bit support and exports")
labeled("What", "A non-modal dialog for per-channel brightness/contrast (display window), gamma, and visibility, with a live per-channel histogram; Fiji-style per-channel pseudo-colour assignment and additive compositing; a publication composite PNG export; and true 16-bit support: a native-bit-depth histogram and native window units computed server-side from the raw pixels (the 8-bit display tiles cannot carry 16-bit values), plus a data-preserving multi-band 16-bit TIFF export of the visible channels. A non-blocking activity indicator and a precomputed recolor LUT keep the controls responsive on large stacks.")
labeled("Where", "libs/jax-image-visualization/src/lib/channel-histogram/*, contracts/channel-histogram-api.contract.ts, and the per-channel/16-bit paths in the OpenSeadragon and Plotly services. Backed by jit-service /histogram and /export/tiff.")
labeled("Status", "Implemented (frontend). Backed by deployed jit-service endpoints.")

h2("D6. Browser SAM segmentation toolset (client-side, ONNX / WebGPU + WASM)")
labeled("What", "In-browser, SAM-based segmentation tools for 2D datasets, running quantized models client-side via onnxruntime-web (WebGPU with a WASM fallback; some models pinned to WASM). Three tools: a promptable box tool that turns drawn rectangles into masks via a real Segment-Anything model (issue #90), an interactive point tool for click-to-segment refinement with a model picker, and an automatic Cellpose-SAM tool. Cellpose-SAM is not itself a promptable model — it pairs SAM's ViT image encoder with a Cellpose flow head and drops SAM's prompt encoder/mask decoder, so it cannot take a box or point directly. We make it effectively box-promptable with a browser-based slide-cropper engine (toolbar/crop/slide-crop.ts), the client-side equivalent of the jit-service slide-crop: it crops the user-drawn rectangle out of the loaded image frame and feeds only that region to Cellpose-SAM, then offsets the returned cell masks back onto the full frame — so the user's box scopes the automatic segmentation to an area of interest. The promptable SAM models are pluggable through a registry that ships hosted ONNX pairs as defaults — micro-sam ViT-T and ViT-B (microscopy-finetuned, promptable) and patho-sam ViT-B (fp16 and int8). The architecture (two-stage encoder/decoder, model registry, tool wiring) and the ONNX export + quantization recipe are documented in docs/sam-segmentation-design.md.")
labeled("Where", "libs/jax-image-visualization/src/lib/toolbar/segmentation/* (sam-tool, sam-point-tool, cell-segment-tool, cellpose-segmenter, sam-model-registry, the ONNX session + worker), the browser slide-cropper at toolbar/crop/slide-crop.ts (crops the prompt rectangle that feeds Cellpose-SAM), and the Segment toolbar buttons. The registry ships hosted ONNX defaults, so no host wiring is required; setSamModelUrls() can repoint a model.")
labeled("Status", "Implemented (2D): box-prompt, point-prompt and automatic Cellpose-SAM are wired, unit-tested, and run client-side against hosted ONNX models shipped as registry defaults (micro-sam ViT-T / ViT-B, patho-sam ViT-B + int8) — no host wiring required.")
labeled("Future work", "Additional fine-tuned Cellpose-SAM variants and further promptable models (e.g. SAM3) as their quantized ONNX exports are published, plus 3D / z-stack mask propagation with cross-slice linking. See §9.")

h2("D7. Extraction into the jax-image-visualization Nx buildable library")
labeled("What", "Move the visualization layer into a standalone Nx buildable library with a clean public boundary: library-owned data interfaces (IImageInfo / IImageMetadata) so the package never imports the host app's concrete models; contracts and dependency-injection ports/tokens (CHANNEL_HISTOGRAM_API, the region-IO port, the tile-access port, and a VIZ_CONFIG injection token) so the host supplies app-specific behaviour; a barrel public API (src/index.ts); an ng-packagr build target; and Jest unit tests. The host app (jit-ui) provides the ports, binding them to its concrete services via useExisting.")
labeled("Where", "libs/jax-image-visualization/ (src/index.ts barrel, src/lib/contracts/ports/*, ng-package.json / project.json build target, *.spec.ts), and apps/jit-ui app module providers.")
labeled("Status", "In progress — the library is buildable (ng-packagr), tested (Jest), and consumed by jit-ui through DI ports; remaining work is auditing for any residual host coupling and finalizing the public API surface.")
labeled("Examples + example server", "Bundle a set of runnable usage examples in the library, powered by a small example server. The example server stands in for jit-service by serving the sample data the viewer needs (preview / tiles-info / tile, histogram, and zoom/region crops), so the examples run standalone from the library repo with no host app and no production backend. The examples demonstrate the main integration paths — embedding <visualization>, the Channels and Histogram dialog, and the region editor, plus implementing the DI ports against the example server — and are shipped with the package (an examples/ app + the lightweight server) as a living reference for downstream adopters.")

h2("D8. Distribution and downstream adoption")
labeled("What", "Publish the library to npm under the @jax-data-science organization scope as an Nx publishable package — candidate names @jax-data-science/image-visualization, @jax-data-science/jax-image-visualization, or @jax-data-science/jit-image-visualization. Establish a versioning and breaking-change policy, document the public API and theming/customization surface, remove any remaining host coupling found in D7, and prove reuse by consuming the published package from a second JAX Data Science frontend.")
labeled("Where", "libs/jax-image-visualization packaging config + a downstream consumer repo (TBD).")
labeled("Status", "Planned.")

# ---------------- 5. PHASING ----------------
h1("5. Phasing and acceptance criteria")
para("Phases 0-4 are implemented and merged to master; Phase 5 is in progress and Phase 6 is planned.")

h2("5.1 Phase 0 — Contract and router (Implemented)")
labeled("Acceptance", "Image renders through OpenSeadragon and the other plot types through Plotly, behind one IVisualizer contract; region state and display options stay in sync across a backend switch.")

h2("5.2 Phase 1 — OpenSeadragon image backend (Implemented)")
labeled("Acceptance", "A whole-slide / z-stack image tiles, zooms, and scrubs through OpenSeadragon with a navigator and scale bar; cached slices scrub without re-fetch.")

h2("5.3 Phase 2 — Regions and tools (Implemented)")
labeled("Acceptance", "Create/edit/select/delete, the wand, vertex eraser and zoom-to-box, and QuPath GeoJSON round-trip all work on both backends from the shared region store.")

h2("5.4 Phase 3 — Plot-type framework (Implemented)")
labeled("Acceptance", "The plot-type selector switches types with correct capability gating; isosurface and intensity-profile inset render for grayscale stacks.")

h2("5.5 Phase 4 — Channels and Histogram + 16-bit (Implemented)")
labeled("Acceptance", "Per-channel window/gamma/pseudo-colour recolor live; composite PNG export works; on a 16-bit stack the histogram and window read in native units and the 16-bit TIFF export downloads a data-preserving multi-band file.")

h2("5.6 Phase 5 — Library extraction hardening (In progress)")
labeled("Goal", "Confirm the package carries no host coupling and freeze the public API.")
labeled("Acceptance", "ng-packagr build is green, Jest tests pass, and a dependency-boundary check confirms the library imports nothing from the host app; the barrel exports the full intended public surface.")
labeled("Task — examples + example server", "Add an example server and a bundled set of usage examples (D7). The example server serves sample preview/tile/histogram/zoom-region data so the examples run standalone, exercising the same tile/preview/zoom paths the production viewer uses. Acceptance: the examples run from the library repo against the example server with no host app, and cover embedding <visualization>, the Channels and Histogram dialog, the region editor, and implementing the DI ports.")

h2("5.7 Phase 6 — Distribution + downstream consumer (Planned)")
labeled("Goal", "Publish the library and prove reuse.")
labeled("Acceptance", "The library is consumable as a versioned artifact and is rendered in a second frontend with no source copy.")

h2("5.8 Test coverage")
para(
    "The library carries a Jest unit-test suite — 682 tests across 43 suites, all passing — run in CI with "
    "coverage reporting and enforced minimum thresholds, so the build fails if coverage regresses below the "
    "gate. Current coverage on master:"
)
table(
    [
        ["Lines", "73.2%", "5198 / 7098", "69%"],
        ["Statements", "71.2%", "6004 / 8436", "67%"],
        ["Functions", "63.4%", "974 / 1536", "60%"],
        ["Branches", "55.8%", "1768 / 3168", "51%"],
    ],
    ["Metric", "Covered", "Covered / total", "CI gate (min)"],
)
para(
    "Coverage concentrates on the backend-neutral logic the library owns: the region store and the on-canvas "
    "tools (wand, brush, vertex-eraser, zoom-to-box), coordinate transforms, the OpenSeadragon tiled load and "
    "slice cache, the display/recolor pipeline, the channel-histogram and region-editor components, and the "
    "Plotly viewport/state methods. The lower branch figure reflects defensive backend/DOM guards that are "
    "harder to exercise headlessly rather than untested features."
)

# ---------------- 6. DEPENDENCIES ----------------
h1("6. Dependencies")
bullet("jit-service tile/histogram/export endpoints (/tiles/info, /tile, /histogram, /export/tiff), including the per-channel tiles and the 16-bit histogram + TIFF export. These are already deployed — the histogram and export endpoints are in place, so no jit-service redeploy is required.")
bullet("Angular 17 + Nx 16 build toolchain and ng-packagr for the buildable library; Jest for unit tests.")
bullet("Third-party rendering libraries: OpenSeadragon (tiled viewer), Plotly (scientific plots), image-js, file-saver.")
bullet("Distribution channel decision (D8) depends on the JAX Data Science platform's package-hosting convention.")

# ---------------- 7. RISKS ----------------
h1("7. Risks")
table([
    ["Residual host coupling blocks a clean extraction",
     "Medium", "Medium",
     "Enforce a dependency-boundary lint rule; route every app-specific behaviour through a DI port; library-owned data interfaces already replace host models."],
    ["Buildable-library AOT / ng-packagr constraints (NG3001, partial-compilation gotchas)",
     "Medium", "Low",
     "Keep the public API in the barrel; build the library in CI on every change; these were already hit and resolved during extraction."],
    ["OpenSeadragon tiling/recolor performance on large grayscale stacks",
     "Medium", "Medium",
     "Precomputed recolor LUT (done); scope per-window-change invalidation to the visible slice if needed; bound tile counts for per-channel rendering."],
    ["16-bit display fidelity is limited because viewer tiles are 8-bit",
     "Low", "Low",
     "True precision is preserved in the 16-bit TIFF export and the native histogram; the display approximation is documented and does not alter the data."],
    ["Versioning / breaking changes destabilize downstream consumers",
     "Medium", "Medium",
     "Semantic versioning + a documented breaking-change policy as part of D8; keep the public surface small and contract-first."],
], header=["Risk", "Likelihood", "Impact", "Mitigation"])

para("The paragraphs below expand each row.", italic=True)
labeled("Residual host coupling (Medium / Medium)",
        "The library must import nothing from jit-ui, but a stray import, hardcoded path or environment reference "
        "can survive the move and still compile inside the monorepo — only breaking when the package is consumed "
        "elsewhere. The dependency-boundary lint rule turns such a leak into a CI failure; the DI ports keep every "
        "app-specific behaviour injected rather than imported; and library-owned data interfaces (IImageInfo / "
        "IImageMetadata) replace the host's concrete models. Boundary risk, mitigated by mechanical enforcement.")
labeled("ng-packagr / AOT constraints (Medium / Low)",
        "Angular library builds (ng-packagr, partial / Ivy compilation) are stricter than app builds; NG3001 fires "
        "when the public API references a symbol that is not exported from the barrel — code that compiles as part "
        "of the app but fails the library target. Impact is low because these are deterministic, well-understood "
        "errors with known fixes, and were already hit and resolved during extraction; keeping the public surface in "
        "src/index.ts and building the library in CI on every change catches regressions early.")
labeled("OpenSeadragon tiling / recolor performance (Medium / Medium)",
        "The client-side recolor pipeline (grayscale colormap LUT plus per-channel window/gamma) runs across "
        "tiles × channels × slices, and a naive display change could invalidate far more than the visible region "
        "— worst on the large multi-channel stacks the tool exists for. It degrades responsiveness, not data. A "
        "precomputed recolor LUT (done) turns recolouring into a lookup; per-window-change invalidation can be scoped "
        "to the visible slice; and per-channel tile counts are bounded to cap memory and work.")
labeled("16-bit display fidelity (Low / Low)",
        "Canvas and the viewer's tiles are 8-bit, so 16-bit images are shown as an 8-bit approximation — a concern "
        "only if mistaken for full precision. It is inherent and bounded, and never alters the data: the "
        "native-bit-depth histogram and windowing read true 16-bit values (computed server-side, since 8-bit tiles "
        "cannot carry them) and the multi-band 16-bit TIFF export preserves them. The display approximation is documented.")
labeled("Versioning / breaking changes (Medium / Medium)",
        "Once published (D8) and consumed by a second frontend, any change to the public API can break those "
        "consumers — and the larger the exposed surface, the more ways to break them. Semantic versioning plus a "
        "documented breaking-change policy let consumers pin versions and anticipate majors, and keeping the surface "
        "small and contract-first (IVisualizer + the DI tokens) minimises both the chance and the blast radius of a break.")

# ---------------- 8. OPEN QUESTIONS ----------------
h1("8. Open questions")
num("npm publishing under @jax-data-science — confirm public npm vs a private/internal registry, plus the Nx publishable-library release setup.")
num("Package name within the @jax-data-science scope — image-visualization, jax-image-visualization, or jit-image-visualization (i.e. @jax-data-science/<name>)?")
num("Theming / customization API — how do downstream consumers override colours, tools, and the tile-source backend?")
num("Which residual components (if any) are too jit-ui-coupled to move yet, and what is the deferral plan?")
num("First external consumer — jit-ui already consumes the library in-repo (the main server-backed viewer and the serverless pipeline preview), but it is the origin host using the workspace package. The open question is the first separate JAX Data Science frontend to consume the published package — proving cross-app reuse — and on what timeline.")

# ---------------- 9. FUTURE WORK & INVESTIGATIONS ----------------
h1("9. Future work & investigations")
para("Exploratory directions under consideration. These are not committed deliverables and "
     "are out of scope for the work in §4; they are recorded here to capture intent and inform "
     "later planning.", italic=True)
num("Image-visualization MCP server — a Model Context Protocol (MCP) server that exposes the "
    "visualization as a set of tools/resources so an LLM client (e.g. a Claude session) can drive "
    "it programmatically: open an image, switch plot type, pan/zoom to a region, toggle "
    "channels/colormap, create/select/edit regions, run browser-side SAM/Cellpose, and read back "
    "the current view or region geometry (e.g. “zoom to the largest region and export it as "
    "GeoJSON”). Investigated under jit-ui#97: the key finding is that the whole UI is already "
    "driven through one backend-neutral command interface (IVisualizer / RoutingVisualizerService), "
    "so the work is to expose that surface over a transport rather than design a new control API. "
    "The recommended approach is a live-control bridge — a thin, serializable control adapter over "
    "IVisualizer behind an opt-in entry point, a WebSocket the host opens, and a separate Node MCP "
    "server package mapping one tool per command — with an optional headless (Playwright) variant "
    "for automation and the stateless plot-generator approach explicitly declined. Builds on the "
    "stable library contract (D7). Full design: docs/mcp-control-design.md.")
num("More segmentation models — extend the segmentation toolset (D6) with additional fine-tuned "
    "Cellpose-SAM variants for other tissue/cell domains and further promptable models (e.g. SAM3) "
    "as their quantized ONNX exports are published. New models drop into the existing pluggable "
    "registry, so adding one is largely an export-and-host step.")
num("3D / z-stack segmentation — propagate SAM masks across slices with cross-slice linking, so a "
    "2D prompt extends through a volume.")

# ---------------- 10. CHANGELOG ----------------
h1("10. Changelog")
labeled(date.today().isoformat(), "Initial draft. Captures the implemented visualization rework (contract + dual backends, regions/tools, plot-type framework, Channels and Histogram incl. 16-bit) and the in-progress library extraction (D7) and planned distribution (D8).")
labeled(date.today().isoformat(), "Distribution decided: npm under the @jax-data-science org scope (Nx publishable). Package name to be chosen among image-visualization / jax-image-visualization / jit-image-visualization. Updated the distribution deliverable and the open questions accordingly.")
labeled(date.today().isoformat(), "Merged to master: the 'openseadragon' branch landed on jit-ui master via PR #79; statuses updated from 'on the openseadragon branch' to 'merged to master'. Renumbered deliverables so the implemented browser SAM segmentation toolset (box/point/Cellpose-SAM) is D6 (extraction → D7, distribution → D8); added future-work items: more fine-tuned Cellpose-SAM / promptable models, 3D propagation, and an image-visualization MCP server (§9).")
labeled(date.today().isoformat(), "Investigated the image-visualization MCP server (jit-ui#97): expanded the §9 future-work item with the investigation outcome (live-control bridge over the existing IVisualizer surface; headless variant; stateless generator declined) and linked the dedicated design doc (docs/mcp-control-design.md) and SOW (JIT_image_visualization_MCP_server_SOW.docx).")
labeled(date.today().isoformat(), "Added §2 'Originality and prior art' — a competitive-landscape comparison against desktop tools (Fiji, napari, QuPath, Imaris, 3D Slicer, ITK-SNAP), server-backed web platforms (OMERO, Cytomine, CaMicroscope) and embeddable JS libraries (OpenSeadragon, Plotly, Annotorious, Viv, Vitessce, Cornerstone3D/OHIF). States precisely what is original (unifying WSI tiling + scientific plots + GeoJSON annotation + true 16-bit + in-browser ViT segmentation in one embeddable, backend-neutral library; the port/adapter architecture; client-side domain-finetuned micro-sam/patho-sam/Cellpose-SAM) and what is not (in-browser SAM as a technique, GPU multichannel rendering where Viv leads, and deep 3D/volumetric and whole-slide analysis where desktop tools remain more capable). Renumbered §3–§11.")

# ---------------- 11. CURRENT VISUALS ----------------
h1("11. Current visuals")


def visual(images, caption):
    """Embed one or more screenshots (centered, width-fit) followed by an italic caption."""
    for img in images:
        doc.add_picture(os.path.join(VIS_DIR, img), width=VIS_WIDTH)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption).italic = True


visual(
    ["fig3-multichannel-fluorescence.png"],
    "Figure 3 — Multi fluorescence image with 4 channels.",
)
visual(
    ["fig4-region-polygon.png", "fig4-region-bezier.png", "fig4-region-classes.png"],
    "Figure 4 — Region overlay with Bézier curve/Polygon shapes",
)
visual(
    ["fig5-openseadragon-tiling.png"],
    "Figure 5 — Large image visualization using OpenSeaDragon & tiling",
)
visual(
    ["fig6-ct-stack.png", "fig6-ct-isosurface.png"],
    "Figure 6 — CT 3d dataset (nii), as an image stack and an isosurface",
)
visual(
    ["fig7-plotly-contour.png"],
    "Figure 7 — Plotly Contour plotting",
)

doc.save(OUT_PATH)
print(f"Wrote {OUT_PATH}")
